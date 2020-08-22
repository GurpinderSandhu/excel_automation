from openpyxl import load_workbook
from openpyxl import utils

import docx
from docx import Document

from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.section import WD_ORIENT

from docx.shared import Pt
from docx.shared import Cm, Inches

from os import listdir
import time
import re
import string
from pathlib import Path
import subprocess
import os
#later implement sys arg excel file name

#-----------------------------BASIC METHODOLOGY----------------------------#
'''
-gather title of interest's coordinates 
-test plan methodology objects 
-table objects
-testId object holds all tables from test methodology tab
-from start find end , i.e have top left corner and bottom right corner of table
-use these coordinates to recieve the data from sheet
-use microsoft word template "i.e template4.docx" that contains pre-formatted custom table style
-only then can you use this table style
-construct document using this table style and adding sections where need be
'''
#--------------------------------------------------------------------------#

#-----------------------------CURRENT ASSUMPTIONS--------------------------#
'''
-only one excel file in this directory
-Assuming table widths (i.e # of columns) to be fixed amount
-Titles of tables are Merged cells beginnning in column A
-Data is directly under the table title
-First column must be populated for each row we assume to be populated (i.e blank rows will terminate the table)
-tpmtable and mytable are essentially the same, do not remember why i did this seperate
'''
#--------------------------------------------------------------------------#
#-----------------------------HARDCODED VALUES-----------------------------#
SD_cols = 6
PR_cols = 3
TM_cols = 11
PI_cols = 6
HO_cols = 4
#tpm = 10 data columns, specifiy length of each column
tpm_widths = [Inches(0.5),Inches(0.5),Inches(3),Inches(1),Inches(1),Inches(1),Inches(0.5),Inches(1),Inches(0.5),Inches(1),Inches(1)]
sd_widths = [Inches(0.5), Inches(1), Inches(4.5),Inches(3),Inches(0.5),Inches(1.5)]

template_path = Path("./dependencies/template4.docx")
testplan_template_path = Path("./dependencies/Test Plan Word Doc Template.docx")

CCFI_section_key_reg = r"[][]"
date_reg = r"\d\d\d\d-\d\d-\d\d"
#--------------------------------------------------------------------------#

#---------------------------TABLE CLASSES----------------------------------#
class myTable:
    def __init__(self,title,coords,begin,end,rows,cols,style,headers_style,data,sheet):
        self.title = title
        self.coords = coords
        self.begin = begin
        self.end = end
        self.rows = rows
        self.cols = cols
        self.style = style
        self.headers_style = headers_style
        self.data = data
        self.sheet = sheet

class testID:
    def __init__(self,title,ID,coords,desc,SD_begin,SD_end,PR_begin,PR_end,TM_begin,TM_end,PI_begin,PI_end,style,headers_style,data):
        self.title = title
        self.ID = ID
        self.coords = coords
        self.desc = desc
        self.SD_begin = SD_begin
        self.SD_end = SD_end
        self.PR_begin = PR_begin
        self.PR_end = PR_end
        self.TM_begin = TM_begin
        self.TM_end = TM_end
        self.PI_begin = PI_begin
        self.PI_end = PI_end
        self.style = style
        self.headers_style = headers_style
        self.data = data
        self.tables = []
    
    def make(self):
        wordTable = doc.add_table(rows = self.rows, cols = self.cols)
        wordTable.style = self.style

class tpmTable:
    def __init__(self,title,ID,begin,end,rows,cols,style,headers_style,data):
        self.title = title
        self.ID = ID
        self.begin = begin
        self.end = end
        self.rows = rows
        self.cols = cols
        self.style = style
        self.headers_style = headers_style
        self.data = data
#--------------------------------------------------------------------------#

#initializing table objects (except for test plan methodology)
TPH = myTable("Test Plan History", 0,0,0,0,0,None,None,0,"Test Plan History")
EC = myTable("Error Codes", 0,0,0,0,0,None,None,0,"Error Codes")
CuR = myTable("Customer Requirements", 0,0,0,0,2,None,None,0,"Test Requirements")
CeR = myTable("Requirements", 0,0,0,0,2,None,None,0,"Test Requirements")
CTFR = myTable("Customer Test Firmware Requirements", 0,0,0,0,2,None,None,0,"Customer Firmware Requirements")
HO = myTable("Hardware Overview", 0,0,0,0,HO_cols,None,None,0,"Hardware Overview")

#grouping the tables
tables = [TPH, EC, CuR, CeR, CTFR, HO]
test_ids = []

cwd = os.getcwd()
suffix = ".xlsx"
filenames = listdir(cwd)
excel_file = [filename for filename in filenames if filename.endswith(suffix)]

#doc is global var used to work with Word doc
doc = Document(template_path)

#assuming only one excel file in folder
wb = load_workbook(filename = excel_file[0])
ws = wb['Test Plan History']

sheet_min_row = ws.min_row
sheet_max_row = ws.max_row
sheet_min_col = ws.min_column
sheet_max_col = ws.max_column

def checkSheets():
    global tables
    non_existent = []
    for table in tables:
        try:
            currentWorksheet(table.sheet)
        except:
            non_existent.append(table)
    tables = [table for table in tables if table not in non_existent]
    print(tables)

def currentWorksheet(sheet_name):
    global ws
    global wb
    ws = wb[sheet_name]

def findTitles(sheet_name):
    currentWorksheet(sheet_name)
    for row in ws.iter_rows(min_row= 1, max_col=13):
        for cell in row:
            if cell.value == "Test Plan History":
                TPH.coords = cell.coordinate
            if cell.value == "Error Codes":
                EC.coords = cell.coordinate
            if cell.value == "Customer Requirements":
                CuR.coords = cell.coordinate 
            if cell.value == "Requirements":
                CeR.coords = cell.coordinate
            if cell.value == "Customer Test Firmware Requirements":
                CTFR.coords = cell.coordinate
            if cell.value == "Hardware Overview":
                HO.coords = cell.coordinate

def findTestIDs(sheet_name):
    currentWorksheet(sheet_name)
    for row in ws.iter_rows(min_row= 1, max_col=13):
        for index,cell in enumerate(row):
            if cell.value == "Test ID:":
                number = row[index+1].value #add test id to string to make unique
                
                key = number + row[index+2].value
                title = cell.value + " " + number + " " +row[index+2].value
            
            if cell.value == "Test Description":
                desc = row[index+2].value
            
                test_ids.append(testID(title,number,cell.coordinate,desc,0,0,0,0,0,0,0,0,None,None,0))

def findRanges():
    currentWorksheet("Test Plan Methodology")
    for i,obj in enumerate(test_ids):
        try:
            max_row = test_ids[i+1].SD_begin[0]
        except:
            max_row = ws.max_row
        for row in ws.iter_rows(min_row= obj.SD_begin[0], max_row=max_row, min_col=1,max_col=14):
            for cell in row:
                if cell.value == "Pre-Requisites":
                    obj.PR_begin = title_to_begin(cell.coordinate)
                    obj.tables.append(tpmTable("Pre-Requisites",obj.ID,obj.PR_begin,None,None,PR_cols,None,None,None))
                if cell.value == "Test Methodology":
                    obj.TM_begin = title_to_begin(cell.coordinate)
                    obj.tables.append(tpmTable("Test Methodology",obj.ID,obj.TM_begin,None,None,TM_cols,None,None,None))

def title_to_begin(coord):
    tmp = utils.cell.coordinate_to_tuple(coord)
    return((tmp[0]+1,tmp[1]))

def findBegin():
    for i in tables:
        if i.coords != 0:
            tmp = utils.cell.coordinate_to_tuple(i.coords)
            i.begin = (tmp[0]+1,1)
        else:
            continue
    
    for i in test_ids:
        if i.coords != 0:
            tmp = utils.cell.coordinate_to_tuple(i.coords)
            i.SD_begin = (tmp[0]+3,1)
            i.tables.append(tpmTable("Supporting Documentation",i.ID,i.SD_begin,None,None,SD_cols,None,None,None))
        else:
            continue

def findAmountCols(val):
    row = ws[val]
    count=0
    cols = 0
    for i in row:
        if i.value == None:
            cols = count
            break
        count+=1
    return(cols)

def findEndTestRequirements():
    currentWorksheet("Test Requirements")
    tmp_begin = [CuR.begin,CeR.begin]
    ends= []

    for i in tmp_begin:
        tmp_value = "tmp"
        count = 0
        while tmp_value != None:
            tmp_value = ws.cell(row = i[0]+count, column = i[1]).value
            if tmp_value == None:
                break
            count+=1
        ends.append(i[0]+count-1)
    count = 0
    for i in tables:
        if i.sheet == "Test Requirements":
            i.end = (ends[count],i.cols)
            count+=1

def findEnd(sheet_name):
    currentWorksheet(sheet_name)
    tmp_begin = [i.begin for i in tables if i.sheet == sheet_name]

    final_ends = []
    ends= []

    for i in tmp_begin:
        tmp_value = "tmp"
        count = 0
        while tmp_value != None:
            tmp_value = ws.cell(row = i[0]+count, column = i[1]).value
            if tmp_value == None:
                break
            count+=1
        ends.append((i[0]+count-1,i[1]))

    for i in ends:
        if i == 0:
            continue
        tmp = findAmountCols(i[0])
        if sheet_name == "Hardware Overview":
            tmp = [i.cols for i in tables if i.sheet == sheet_name][0]        
        final_ends.append((i[0],tmp))

    count = 0
    
    current_table = [i for i in tables if i.sheet == sheet_name]
    current_table[0].end = final_ends[0]
 
def TPMfindEnd(tbls):
    currentWorksheet('Test Plan Methodology')
    for tbl in tbls:
        end_col = (tbl.begin[1] + tbl.cols)-1
        tmp = "tmp"
        count = 0
        while tmp != None:
            tmp = ws.cell(row = tbl.begin[0]+count, column = tbl.begin[1]).value
            if tmp == None:
                break
            count+=1
        tbl.end = (tbl.begin[0]+count-1,end_col)
        tbl.rows = count

def getData(range,sheet_name):
    tmp = []
    currentWorksheet(sheet_name)
    print(sheet_name, range)
    for row in ws.iter_rows(min_row= range[0][0], max_row=range[1][0], min_col=range[0][1],max_col=range[1][1]):
        for cell in row:
            if regexCheck(date_reg,str(cell.value)) and not sheet_name == "Test Plan Methodology":
                print(cell.value)
                tmp.append(str(cell.value.date()))
            else:
                tmp.append(str(cell.value))
    return(tmp)

def regexCheck(regex, value):
    if re.search(regex,str(value)):
        return(True)
    else:
        return(False)

def setRowsCols(table):
    if table.begin == 0 or table.end == 0:
        table.rows = 0
        table.cols = 0
        return
    table.rows = table.end[0] - table.begin[0] + 1
    table.cols = table.end[1] - table.begin[1] + 1

def tpmConstructor():
    for i in test_ids:
        TPMfindEnd(i.tables)
        for j in i.tables:
            j.data = getData((j.begin,j.end),"Test Plan Methodology")

def dataConstructor():
    for i in tables:
        setRowsCols(i)
        i.data = getData((i.begin,i.end),i.sheet)

def makeTable(table):
    dtable = doc.add_table(rows = table.rows, cols = table.cols)
    table.style = doc.styles['Colour']
    dtable.style = table.style
    dtable.cell(0,0).vertical_alignment = WD_ALIGN_VERTICAL.TOP

    data_count = 0

    for i, row in enumerate(dtable.rows):
        if i == 0:
            title_row = True
        for cell in row.cells:
            delete_paragraph(cell.paragraphs[0])
            p = cell.add_paragraph(table.data[data_count],style='Body Text')
            p.style.font.name = 'Arial'
            p.style.font.size = Pt(10) 

            if (data_count+1) % table.cols == 0:
                cell.width = Cm(15)
            data_count+=1
        if table.title == "Test Plan History":
            row.cells[0].width = Inches(0.5)

def docTestIDbuilder():
    count=0
    for i in test_ids:
        if count !=0:
            doc.add_section()
        doc.add_paragraph(i.title, style = 'Heading 2')
        doc.add_paragraph("Test Description", style = 'Heading 3')
        if i.desc == None:
            i.desc = ""
        doc.add_paragraph("\t"+i.desc, style = 'Body Text')
        docTestIDtables(i)
        count+=1

def docTestIDtables(test_id):
    for tbl in test_id.tables:
        tm_flag = False
        sd_flag = False
        if tbl.title == "Test Methodology":
            tm_flag = True
        if tbl.title == "Supporting Documentation":
            sd_flag = True
        
        doc.add_paragraph(tbl.title, style = 'Heading 3')
        dtable = doc.add_table(rows = tbl.rows, cols = tbl.cols)
        dtable.style = doc.styles['Colour']

        if tm_flag:
            dtable.allow_autofit = False

        data_count = 0
        for row, i in enumerate(dtable.rows):
            for column, cell in enumerate(i.cells):
                if row == 0:
                    if column == 6:
                        cell.merge(i.cells[column+1])
                    if column == 8:
                        cell.merge(i.cells[column+1])
                    if column == 11:
                        delete_paragraph(cell.paragraphs[0])
                        p = cell.add_paragraph("Error Code",style='Body Text')
                        data_count+=1
                        continue
                    if column == 12:
                        delete_paragraph(cell.paragraphs[0])
                        p = cell.add_paragraph("Test Req Status",style='Body Text')
                        continue

                delete_paragraph(cell.paragraphs[0])
                p = cell.add_paragraph(tbl.data[data_count],style='Body Text')
                p.style.font.name = 'Arial'
                p.style.font.size = Pt(10)

                if (data_count+1) % tbl.cols == 0 and not tm_flag:
                    cell.width = Cm(15)
                if tm_flag and column < 11:
                    cell.width = tpm_widths[column]
                if sd_flag:
                    cell.width = sd_widths[column]
 
                if column == 0:
                    cell.width = Inches(0.5)
                data_count+=1

def docBuilder():
    #flag signals if no CTFR or HO
    flag = False

    doc.add_paragraph("Test Plan History", style = 'Heading 1')
    makeTable(TPH)
    doc.add_section()

    doc.add_paragraph("Customer Requirements", style = 'Heading 1')
    makeTable(CuR)
    doc.add_section()

    doc.add_paragraph("Requirements", style = 'Heading 1')
    makeTable(CeR)
    doc.add_section()

    if(CTFR in tables):
        doc.add_paragraph("Firmware Requirements", style = 'Heading 1')
        makeTable(CTFR)
        doc.add_section()
        flag = True

    if(HO in tables):
        doc.add_paragraph("Hardware Overview", style = 'Heading 1')
        makeTable(HO)
        doc.add_section()
        flag = True

    doc.add_paragraph("Test Plan Methodology", style = 'Heading 1')
    docTestIDbuilder()
    doc.add_section()

    doc.add_paragraph("Error Codes", style = 'Heading 1')
    makeTable(EC)

    if flag:
        portrait = [doc.sections[0],doc.sections[1],doc.sections[2],doc.sections[3],doc.sections[-1]]
    else:
        portrait = [doc.sections[0],doc.sections[1],doc.sections[2],doc.sections[-1]]
  
    landscape = [i for i in doc.sections if i not in portrait]

    for section in landscape:
        new_width, new_height = section.page_height, section.page_width
        section.orientation = WD_ORIENT.LANDSCAPE
        section.page_width = new_width
        section.page_height = new_height
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)
    for section in portrait:
        new_width, new_height = section.page_height, section.page_width
        section.orientation = WD_ORIENT.PORTRAIT
        section.page_width = new_width
        section.page_height = new_height
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

def delete_paragraph(paragraph):
    p = paragraph._element
    p.getparent().remove(p)
    p._p = p._element = None

def main():
    print("Generating Test Plan from Excel document, Please wait .....")

    checkSheets()
    for table in tables:
        findTitles(table.sheet)

    findTestIDs("Test Plan Methodology")
    findBegin()
    findRanges()

    for table in tables:
        if table.sheet == "Test Requirements":
            continue
        findEnd(table.sheet)

    findEndTestRequirements()

    tpmConstructor()
    dataConstructor()
    docBuilder()

    doc.save('tmp.docx')

    #have to import 'Style1' from template3.docx into Test Plan Word Doc Template.docx

    doc1 = docx.Document(testplan_template_path)
    doc2 = docx.Document('tmp.docx')
    for element in doc2.element.body:
        doc1 .element.body.append(element)
    doc1.save('generated_test_plan.docx')
    os.remove('tmp.docx')

if __name__ == '__main__':
    main()